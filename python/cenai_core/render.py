from typing import Sequence

from bs4 import BeautifulSoup
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from weasyprint import HTML


def html_to_pdf(htmls: Sequence[str],
                pdf_file: Path
                ) -> None:

    pdfs = [HTML(string=html).render() for html in htmls]

    pages = []
    for pdf in pdfs:
        pages.extend(pdf.pages)

    document = pdfs[0].copy()
    document.pages = pages
    document.write_pdf(str(pdf_file))


def html_to_docx(htmls: Sequence[str],
                 docx_file: Path
                 ) -> None:
    document = docx.Document()

    for html in htmls:
        soup = BeautifulSoup(html, "html.parser")

        for element in soup.find_all(["h2", "table"]):
            if element.name == "h2":
                heading = document.add_heading(level=2)
                heading.text = element.get_text(strip=True)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

            elif element.name == "table":
                _add_table_to_document(document, str(element))

            document.add_paragraph()

    document.save(docx_file)


def _add_table_to_document(document: docx.Document,
                           table: str):
    soup = BeautifulSoup(table, "html.parser")
    rows = soup.find_all("tr")

    max_cols = max(len(row.find_all(["td", "th"])) for row in rows)
    docx_table = document.add_table(rows=0, cols=max_cols)
    docx_table.style = "Table Grid"

    rowspan_map = [0] * max_cols

    for html_row in rows:
        docx_row = docx_table.add_row()
        col_index = 0

        for html_cell in html_row.find_all(["td", "th"]):
            while col_index < len(rowspan_map) and rowspan_map[col_index] > 0:
                rowspan_map[col_index] -= 1
                col_index += 1

            if col_index < len(docx_row.cells):
                docx_cell = docx_row.cells[col_index]
                docx_cell.text = html_cell.get_text(strip=True)

            if html_cell.has_attr("rowspan"):
                rowspan = int(html_cell["rowspan"])
                for i in range(rowspan):
                    if col_index + i < len(rowspan_map):
                        rowspan_map[col_index + i] = rowspan - 1

            if html_cell.has_attr("colspan"):
                colspan = int(html_cell["colspan"])
                for i in range(1, colspan):
                    if col_index + i < len(docx_row.cells):
                        docx_row.cells[col_index].merge(docx_row.cells[col_index + i])
                col_index += colspan - 1

            col_index += 1
